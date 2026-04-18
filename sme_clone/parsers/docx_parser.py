"""
docx_parser.py — Extract text from Word documents (.docx) into markdown.

This parser:
  1. Reads the .docx file.
  2. Extracts paragraphs, headings, and basic formatting.
  3. Outputs one .md file per document into parsed/docs/.
"""

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def parse_docx(
    docx_path: Path,
    output_dir: Path,
) -> Path:
    """Parse a Word document and write a markdown file.

    Args:
        docx_path: Path to the .docx file.
        output_dir: Directory to write the parsed markdown (e.g., parsed/docs/).

    Returns:
        Path to the generated markdown file.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{docx_path.stem}.md"

    try:
        doc = Document(str(docx_path))
    except PackageNotFoundError:
        output_path.write_text(
            f"# {docx_path.stem}\n\n"
            f"> Parsed from `{docx_path.name}`\n\n"
            "*Error: Could not open this file as a valid .docx document.*\n",
            encoding="utf-8",
        )
        return output_path

    lines: list[str] = []

    # Map Word heading levels to markdown heading levels.
    # The document title becomes h1; Heading 1 → h2, Heading 2 → h3, etc.
    heading_map = {
        "Title": 1,
        "Heading 1": 2,
        "Heading 2": 3,
        "Heading 3": 4,
        "Heading 4": 5,
        "Heading 5": 6,
    }

    doc_title_added = False

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            # Preserve paragraph breaks
            if lines and lines[-1] != "":
                lines.append("")
            continue

        if style_name in heading_map:
            level = heading_map[style_name]
            lines.append(f"{'#' * level} {text}")
            lines.append("")
            if level == 1:
                doc_title_added = True
        elif style_name.startswith("Heading"):
            # Catch any other heading styles
            lines.append(f"### {text}")
            lines.append("")
        elif style_name == "List Bullet" or style_name.startswith("List Bullet"):
            lines.append(f"- {text}")
        elif style_name == "List Number" or style_name.startswith("List Number"):
            lines.append(f"1. {text}")
        elif style_name == "Quote" or style_name == "Intense Quote":
            lines.append(f"> {text}")
            lines.append("")
        else:
            lines.append(text)
            lines.append("")

    # If no title heading was found, prepend one from the filename
    if not doc_title_added:
        title = docx_path.stem.replace("_", " ").replace("-", " ").title()
        lines.insert(0, "")
        lines.insert(0, "")
        lines.insert(0, f"> Parsed from `{docx_path.name}`")
        lines.insert(0, "")
        lines.insert(0, f"# {title}")
    else:
        lines.insert(1, f"> Parsed from `{docx_path.name}`")
        lines.insert(2, "")

    # Extract tables
    if doc.tables:
        lines.append("---")
        lines.append("")
        for table_idx, table in enumerate(doc.tables):
            lines.append(f"### Table {table_idx + 1}")
            lines.append("")
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    lines.append("| " + " | ".join("---" for _ in cells) + " |")
            lines.append("")

    if not any(line.strip() for line in lines if not line.startswith(("#", ">"))):
        lines.append("*No text content could be extracted from this document.*")
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path
